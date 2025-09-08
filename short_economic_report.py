#!/usr/bin/env python3
"""
Short economic report - contains all currency rates,
cheapest item of each type and best region for each product.
"""

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from economy import (
    fetch_countries_and_currencies,
    build_currency_rates_map,
    fetch_cheapest_items_from_all_countries,
    fetch_items_by_type
)
from production_analyzer_consolidated import ProductionAnalyzer
from storage import load_historical_data
from datetime import timedelta


def generate_short_economic_report(
    output_dir: str = "reports",
    sections: dict = None
) -> str:
    """
    Generates a short economic report in DOCX format.
    
    Args:
        output_dir: Directory to save the report
        sections: Sections to include (all by default)
    
    Returns:
        Path to the generated file
    """
    
    # Default sections
    if sections is None:
        sections = {
            'currency_rates': True,
            'cheapest_items': True,
            'best_regions': True
        }
    
    print("🚀 Generating short economic report...")
    
    # Load historical data for currency rate comparison
    historical_data = load_historical_data()
    
    # Get economic data
    print("📊 Fetching economic data...")
    eco_countries, currencies_map, currency_codes_map, gold_id = fetch_countries_and_currencies()
    
    if not eco_countries or not currencies_map:
        print("❌ Error: Cannot fetch economic data")
        return None
    
    # Get currency rates
    print("💰 Fetching currency rates...")
    currency_rates = build_currency_rates_map(currencies_map, gold_id)
    
    # Get items map
    print("📦 Fetching items map...")
    items_map = fetch_items_by_type("economic")
    
    # Get cheapest items
    print("🛒 Fetching cheapest items...")
    cheapest_items = fetch_cheapest_items_from_all_countries(
        eco_countries, items_map, currency_rates, gold_id
    )
    
    # Get region data for productivity analysis
    print("🏭 Fetching region data...")
    from regions import fetch_and_process_regions
    regions_data, regions_summary = fetch_and_process_regions(eco_countries)
    
    # Create DOCX document
    document = Document()
    
    # Title
    title = document.add_heading("Short Economic Report Eclesiar", level=1)
    for run in title.runs:
        run.font.size = Pt(24)
    
    # Date information
    document.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("")
    
    # 1. Currency rates
    if sections.get('currency_rates', True):
        print("📈 Adding currency rates section...")
        document.add_heading("💰 Currency Rates vs GOLD", level=1)
        
        if currency_rates:
            # Get yesterday's rates for comparison
            yesterday_key = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
            yesterday_rates = {}
            
            try:
                if yesterday_key in historical_data:
                    yesterday_rates = (historical_data[yesterday_key].get('economic_summary') or {}).get('currency_rates') or {}
                    print(f"DEBUG: Found yesterday rates: {bool(yesterday_rates)}")
                
                # If no yesterday data, look for the nearest previous day
                if not yesterday_rates and historical_data:
                    available_dates = sorted([d for d in historical_data.keys() if d < yesterday_key], reverse=True)
                    for date in available_dates:
                        rates = (historical_data[date].get('economic_summary') or {}).get('currency_rates') or {}
                        if rates:
                            yesterday_rates = rates
                            print(f"DEBUG: Found rates from {date}: {bool(yesterday_rates)}")
                            break
            except Exception as e:
                print(f"DEBUG: Error getting yesterday rates: {e}")
                yesterday_rates = {}
            
            # Sort rates alphabetically by currency name
            sorted_rates = []
            for currency_id, rate in currency_rates.items():
                currency_name = currencies_map.get(currency_id, f"Currency {currency_id}")
                currency_code = currency_codes_map.get(currency_id, "")
                display_name = currency_code if currency_code else currency_name
                sorted_rates.append((display_name, rate, currency_name, currency_id))
            
            sorted_rates.sort(key=lambda x: x[0].lower())
            
            # Determine table columns based on whether we have historical data
            has_historical = bool(yesterday_rates)
            cols = 5 if has_historical else 3
            
            # Currency rates table
            table = document.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Currency"
            hdr_cells[1].text = "Code"
            hdr_cells[2].text = "Rate (GOLD per 1)"
            if has_historical:
                hdr_cells[3].text = "Previous Rate"
                hdr_cells[4].text = "Change %"
            
            # Data rows
            for display_name, rate, full_name, currency_id in sorted_rates:
                row_cells = table.add_row().cells
                row_cells[0].text = full_name
                row_cells[1].text = display_name
                row_cells[2].text = f"{rate:.6f}"
                
                if has_historical:
                    # Get previous rate for this currency
                    prev_rate = yesterday_rates.get(str(currency_id))
                    if prev_rate is not None:
                        row_cells[3].text = f"{prev_rate:.6f}"
                        # Calculate percentage change
                        if prev_rate > 0:
                            change_pct = ((rate - prev_rate) / prev_rate) * 100
                            row_cells[4].text = f"{change_pct:+.2f}%"
                        else:
                            row_cells[4].text = "N/A"
                    else:
                        row_cells[3].text = "N/A"
                        row_cells[4].text = "N/A"
        else:
            document.add_paragraph("No currency rate data available.")
        
        document.add_paragraph("")
    
    # 2. Cheapest items of each type
    if sections.get('cheapest_items', True):
        print("🛒 Adding cheapest items section...")
        document.add_heading("🛒 Cheapest Items of Each Type", level=1)
        
        if cheapest_items:
            # Group items by type
            grouped_items = _group_items_by_type(cheapest_items)
            
            for item_type, items in grouped_items.items():
                if items:
                    document.add_heading(f"{item_type}", level=2)
                    
                    # Find cheapest item of this type
                    cheapest_item = min(items, key=lambda x: x.get('price_gold', float('inf')))
                    
                    # Table with cheapest item
                    table = document.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Item"
                    hdr_cells[1].text = "Country"
                    hdr_cells[2].text = "Price (currency)"
                    hdr_cells[3].text = "Price (GOLD)"
                    hdr_cells[4].text = "Quantity"
                    hdr_cells[5].text = "Average of 5 cheapest (GOLD)"
                    
                    # Data row
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(cheapest_item.get('item_name', 'Unknown'))
                    row_cells[1].text = str(cheapest_item.get('country', 'Unknown'))
                    row_cells[2].text = f"{cheapest_item.get('price_currency', 0):.3f} {cheapest_item.get('currency_name', '')}"
                    row_cells[3].text = f"{cheapest_item.get('price_gold', 0):.6f}"
                    qty = cheapest_item.get('amount')
                    row_cells[4].text = (str(int(qty)) if isinstance(qty, (int, float)) else "—")
                    avg5 = cheapest_item.get('avg5_in_gold')
                    row_cells[5].text = (f"{float(avg5):.6f}" if isinstance(avg5, (int, float)) else "—")
                    
                    document.add_paragraph("")
        else:
            document.add_paragraph("No cheapest items data available.")
        
        document.add_paragraph("")
    
    # 3. Best regions for each product
    if sections.get('best_regions', True) and regions_data:
        print("🏭 Adding best regions section...")
        document.add_heading("🏭 Best Regions for Production", level=1)
        
        # Analyze productivity
        analyzer = ProductionAnalyzer()
        production_data = analyzer.analyze_all_regions(regions_data)
        
        if production_data:
            # Group by product type
            products_groups = {}
            for data in production_data:
                item_name = data.item_name if hasattr(data, 'item_name') else data.get('item_name', 'Unknown')
                if item_name not in products_groups:
                    products_groups[item_name] = []
                products_groups[item_name].append(data)
            
            # For each product show the best region
            for product_name, items in products_groups.items():
                if items:
                    # Sort by efficiency_score
                    items.sort(
                        key=lambda x: x.efficiency_score if hasattr(x, 'efficiency_score') else x.get('efficiency_score', 0), 
                        reverse=True
                    )
                    
                    # Take the best region
                    best_region = items[0]
                    
                    document.add_heading(f"Product: {product_name.title()}", level=2)
                    
                    # Table with best region
                    table = document.add_table(rows=1, cols=9)
                    table.style = 'Table Grid'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Region"
                    hdr_cells[1].text = "Country"
                    hdr_cells[2].text = "Score"
                    hdr_cells[3].text = "Bonus"
                    hdr_cells[4].text = "Q1"
                    hdr_cells[5].text = "Q2"
                    hdr_cells[6].text = "Q3"
                    hdr_cells[7].text = "Q4"
                    hdr_cells[8].text = "Q5"
                    
                    # Data row
                    row_cells = table.add_row().cells
                    if hasattr(best_region, 'region_name'):
                        row_cells[0].text = best_region.region_name
                        row_cells[1].text = best_region.country_name
                        row_cells[2].text = f"{best_region.efficiency_score:.2f}"
                        row_cells[3].text = f"{best_region.total_bonus:.1%}"
                        row_cells[4].text = str(best_region.production_q1)
                        row_cells[5].text = str(best_region.production_q2)
                        row_cells[6].text = str(best_region.production_q3)
                        row_cells[7].text = str(best_region.production_q4)
                        row_cells[8].text = str(best_region.production_q5)
                    else:
                        row_cells[0].text = best_region.get('region_name', 'Unknown')
                        row_cells[1].text = best_region.get('country_name', 'Unknown')
                        row_cells[2].text = f"{best_region.get('efficiency_score', 0):.2f}"
                        row_cells[3].text = f"{best_region.get('total_bonus', 0):.1%}"
                        row_cells[4].text = str(best_region.get('production_q1', 0))
                        row_cells[5].text = str(best_region.get('production_q2', 0))
                        row_cells[6].text = str(best_region.get('production_q3', 0))
                        row_cells[7].text = str(best_region.get('production_q4', 0))
                        row_cells[8].text = str(best_region.get('production_q5', 0))
                    
                    document.add_paragraph("")
        else:
            document.add_paragraph("No region productivity data available.")
    
    # Save document
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
    filename = f"short_economic_report_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    document.save(filepath)
    print(f"✅ Short economic report generated: {filepath}")
    
    return filepath


def _group_items_by_type(cheapest_items: Dict[int, List[Dict[str, Any]]]) -> Dict[str, List[Dict[str, Any]]]:
    """
    Groups items by type (grain, iron, weapon, etc.)
    """
    import re
    
    def parse_item_type(name: str) -> str:
        """Determines item type based on name"""
        name_lower = str(name).lower()
        
        if any(word in name_lower for word in ["grain", "zboże"]):
            return "Grain"
        elif any(word in name_lower for word in ["iron", "żelazo"]):
            return "Iron"
        elif any(word in name_lower for word in ["titanium", "tytan"]):
            return "Titanium"
        elif any(word in name_lower for word in ["fuel", "paliwo"]):
            return "Fuel"
        elif any(word in name_lower for word in ["food", "żywność"]):
            return "Food"
        elif "weapon" in name_lower or "broń" in name_lower:
            return "Weapon"
        elif "aircraft" in name_lower or ("airplane" in name_lower and "ticket" not in name_lower):
            return "Aircraft"
        elif "ticket" in name_lower:
            return "Airplane Ticket"
        else:
            return "Other"
    
    grouped = {}
    
    for item_id, item_list in cheapest_items.items():
        if item_list and len(item_list) > 0:
            for item in item_list:
                item_name = item.get('item_name', '')
                item_type = parse_item_type(item_name)
                
                if item_type not in grouped:
                    grouped[item_type] = []
                grouped[item_type].append(item)
    
    return grouped


def main():
    """Main function for testing"""
    try:
        report_path = generate_short_economic_report()
        if report_path:
            print(f"✅ Report generated successfully: {report_path}")
        else:
            print("❌ Error generating report")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
