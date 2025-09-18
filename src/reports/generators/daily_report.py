import io
import os
import base64
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt
from src.data.database.models import get_item_price_series, get_item_price_avg


def create_infographic(title, data, labels, color, ax_title):
    if not data or not labels:
        return None
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(10, 6))
    y_pos = np.arange(len(labels))
    ax.bar(y_pos, data, color=color, edgecolor='black', zorder=3)
    ax.set_xticks(y_pos)
    ax.set_xticklabels(labels, rotation=45, ha='right')
    ax.set_title(title, fontsize=16, fontweight='bold', pad=20)
    ax.set_ylabel(ax_title, fontsize=12)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    for i, v in enumerate(data):
        ax.text(i, v + max(data) * 0.05, f'{v:,}', ha='center', fontweight='bold')
    plt.tight_layout()
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=300)
    plt.close(fig)
    buffer.seek(0)
    return buffer





def generate_report(
    summary_data,
    historical_data,
    top_warriors,
    items_map,
    currencies_map,
    country_map=None,
    currency_codes_map=None,
    gold_id=None,
    output_dir="reports",
    sections=None,
):
    """Generate daily report as a DOCX file and return its path."""
    # Set default sections if not provided
    if sections is None:
        sections = {
            'military': True,
            'warriors': True, 
            'economic': True,
            'production': True
        }
    
    document = Document()

    # Title
    title = document.add_heading("Eclesiar's Pulse", level=1)
    for run in title.runs:
        run.font.size = Pt(24)

    document.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    fetched_at = summary_data.get('fetched_at')
    if fetched_at:
        try:
            dt = datetime.fromisoformat(str(fetched_at).replace('Z', '+00:00'))
            ts = dt.astimezone().strftime("%Y-%m-%d %H:%M %Z") if dt.tzinfo else dt.strftime("%Y-%m-%d %H:%M")
            document.add_paragraph(f"Source data fetched: {ts}")
        except Exception:
            document.add_paragraph(f"Source data fetched: {fetched_at}")

    # Best warrior shoutout
    if sections.get('warriors', False):
        tw_local = summary_data.get('top_warriors') or []
        if isinstance(tw_local, list) and len(tw_local) > 0:
            best = tw_local[0]
            best_name = str(best.get('username', 'Unknown Warrior'))
            # Resolve country name to English using provided maps if available
            best_country_raw = best.get('country')
            if not best_country_raw or str(best_country_raw).strip().lower() == 'nieznany kraj':
                cm = country_map or {}
                nat_id = best.get('nationality_id')
                reg_id = best.get('region_id')
                if nat_id in (cm or {}):
                    best_country = cm.get(nat_id)
                elif reg_id in (cm or {}):
                    best_country = cm.get(reg_id)
                else:
                    best_country = 'Unknown Country'
            else:
                best_country = str(best_country_raw)
            p = document.add_paragraph(f"Congratulations to the Top Warrior: {best_name} ({best_country})!")
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(14)

    # Removed welcome blurb per request

    # Military section
    if sections.get('military', False):
        document.add_heading("⚔️ Battlefield Pulse: Key War Statistics", level=1)
        document.add_paragraph("Damage attributed to a country is the sum of damage dealt by players of that country across all battles in the last 24 hours (regardless of whether the country was a party to the conflict).")

    # Define yesterday_key and today's war counts at the beginning
    yesterday_key = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    
    wars_summary = summary_data.get('wars_summary', {}) or {}
    ongoing_wars = wars_summary.get('ongoing') or []
    finished_wars = wars_summary.get('finished') or []
    today_ongoing = len(ongoing_wars)
    today_finished = len(finished_wars)
    
    if sections.get('military', False) and (ongoing_wars or finished_wars):
        document.add_heading("🗺️ Military Activity: Ongoing and Finished Wars", level=2)
        if ongoing_wars:
            document.add_paragraph("Ongoing battles:")
            for line in ongoing_wars:
                # Normalize potential Polish labels from API-provided strings
                txt = str(line)
                txt = txt.replace('wojna', 'war').replace('Wojna', 'War').replace('wynik', 'score').replace('Wynik', 'Score')
                document.add_paragraph(txt, style='List Bullet')
        else:
            document.add_paragraph("No ongoing battles at the moment.")
        if finished_wars:
            document.add_paragraph("Finished battles:")
            for line in finished_wars:
                txt = str(line)
                txt = txt.replace('wojna', 'war').replace('Wojna', 'War').replace('wynik', 'score').replace('Wynik', 'Score')
                document.add_paragraph(txt, style='List Bullet')
        else:
            document.add_paragraph("No finished battles to display.")
        
        # Compare wars with yesterday if available
        if yesterday_key in historical_data and historical_data[yesterday_key].get('wars_summary'):
            yesterday_wars = historical_data[yesterday_key]['wars_summary']
            yesterday_ongoing = len(yesterday_wars.get('ongoing', []))
            yesterday_finished = len(yesterday_wars.get('finished', []))
            
            if today_ongoing != yesterday_ongoing or today_finished != yesterday_finished:
                document.add_heading("🗺️ Wars: Changes vs Yesterday", level=3)
                if today_ongoing != yesterday_ongoing:
                    change = today_ongoing - yesterday_ongoing
                    arrow = "▲" if change > 0 else "▼"
                    document.add_paragraph(f"Ongoing wars: {arrow} {change:+d} ({today_ongoing} vs {yesterday_ongoing})")
                if today_finished != yesterday_finished:
                    change = today_finished - yesterday_finished
                    arrow = "▲" if change > 0 else "▼"
                    document.add_paragraph(f"Finished wars: {arrow} {change:+d} ({today_finished} vs {yesterday_finished})")

    if sections.get('military', False):
        military_summary = summary_data.get('military_summary', {})
        if not military_summary:
            document.add_paragraph("No war data from the last 24 hours. The server was calm.")
        else:
            # Normalize country names (avoid Polish 'Nieznany Kraj')
            def normalize_country_name(name: str) -> str:
                if not name:
                    return 'Unknown Country'
                n = str(name).strip()
                if n.lower() == 'nieznany kraj':
                    return 'Unknown Country'
                return n

            sorted_countries = sorted(military_summary.keys(), key=lambda x: military_summary[x], reverse=True)
            top_countries = sorted_countries[:7]
            document.add_paragraph(f"In the last 24 hours, military activity was recorded in {len(military_summary)} countries.")
            for country in top_countries:
                dmg = military_summary[country]
                document.add_paragraph(f"{normalize_country_name(country)}: {dmg:,}")
        
        # Add executive summary of changes if historical data is available
        if yesterday_key in historical_data:
            document.add_heading("📊 Executive Summary: Key Changes", level=2)
            
            # Military changes summary
            if historical_data[yesterday_key].get('military_summary'):
                yesterday_military = historical_data[yesterday_key]['military_summary']
                total_today = sum(military_summary.values())
                total_yesterday = sum(yesterday_military.values())
                if total_yesterday > 0:
                    military_change_pct = (total_today - total_yesterday) / total_yesterday * 100.0
                    arrow = "▲" if military_change_pct > 0 else ("▼" if military_change_pct < 0 else "→")
                    document.add_paragraph(f"⚔️ Total military activity: {arrow} {military_change_pct:+.1f}% ({total_today:,} vs {total_yesterday:,})")
            
            # Economic changes summary
            economic_data = summary_data.get('economic_summary', {})
            if historical_data[yesterday_key].get('economic_summary'):
                yesterday_economic = historical_data[yesterday_key]['economic_summary']
                today_rates = economic_data.get('currency_rates', {})
                yesterday_rates = yesterday_economic.get('currency_rates', {})
                
                if today_rates and yesterday_rates:
                    # Count currencies with significant changes
                    significant_changes = 0
                    for currency_id in today_rates:
                        if currency_id in yesterday_rates:
                            try:
                                today_val = float(today_rates[currency_id])
                                yesterday_val = float(yesterday_rates[currency_id])
                                if yesterday_val != 0:
                                    change_pct = abs((today_val - yesterday_val) / yesterday_val * 100.0)
                                    if change_pct > 1.0:
                                        significant_changes += 1
                            except (ValueError, TypeError):
                                continue
                    
                    if significant_changes > 0:
                        document.add_paragraph(f"💰 Currency volatility: {significant_changes} currencies changed by >1%")
                    else:
                        document.add_paragraph("💰 Currency markets: Stable (no significant changes)")
            
            document.add_paragraph("See detailed comparisons below for specific metrics.")
        if yesterday_key in historical_data and historical_data[yesterday_key].get('military_summary'):
            document.add_heading("⚖️ Military: Compared to yesterday", level=2)
            yesterday_data = historical_data[yesterday_key]['military_summary']
            
            # Show top 5 countries with biggest changes
            changes = []
            for country in set(list(military_summary.keys()) + list(yesterday_data.keys())):
                today_dmg = military_summary.get(country, 0)
                yesterday_dmg = yesterday_data.get(country, 0)
                change = today_dmg - yesterday_dmg
                if change != 0:  # Only show countries with changes
                    changes.append((abs(change), change, normalize_country_name(country)))
            
            if changes:
                # Sort by absolute change value
                changes.sort(key=lambda x: x[0], reverse=True)
                document.add_paragraph("Top changes in military activity:")
                for _, change, country in changes[:5]:
                    if change > 0:
                        document.add_paragraph(f"• {country}: +{change:,} (▲ increase)")
                    else:
                        document.add_paragraph(f"• {country}: {change:,} (▼ decrease)")
            else:
                document.add_paragraph("No significant changes in military activity detected.")

        # Economic comparison with yesterday has been removed per request

    # Top warriors
    if sections.get('warriors', False):
        document.add_heading("🏆 Best of the Best: Heroes Ranking", level=1)
        top_warriors_local = summary_data.get('top_warriors', [])
        if not top_warriors_local:
            document.add_paragraph("No top warriors data from the last day.")
        else:
            document.add_paragraph("Here's the list of the most powerful warriors from the last 24 hours:")
            cm = country_map or {}
            for warrior in top_warriors_local:
                username = warrior.get('username', f"ID: {warrior.get('id', 'Unknown')}")
                damage = warrior.get('damage', 0)
                country = warrior.get('country')
                if not country or str(country).strip().lower() == 'nieznany kraj':
                    nat_id = warrior.get('nationality_id')
                    reg_id = warrior.get('region_id')
                    if nat_id in cm:
                        country = cm.get(nat_id)
                    elif reg_id in cm:
                        country = cm.get(reg_id)
                    else:
                        country = 'Unknown Country'
                document.add_paragraph(f"{username} ({country}): {damage:,} damage", style='List Number')
            
            # Compare with yesterday's top warriors if available
            if yesterday_key in historical_data and historical_data[yesterday_key].get('top_warriors'):
                yesterday_warriors = historical_data[yesterday_key]['top_warriors']
                if yesterday_warriors:
                    document.add_heading("🏆 Top Warriors: Changes vs Yesterday", level=2)
                    
                    # Find warriors who appeared in both days
                    today_usernames = {w.get('username', '').lower() for w in top_warriors_local}
                    yesterday_usernames = {w.get('username', '').lower() for w in yesterday_warriors}
                    
                    # Warriors who maintained position
                    maintained = today_usernames.intersection(yesterday_usernames)
                    if maintained:
                        document.add_paragraph("Warriors who maintained top positions:")
                        for username in list(maintained)[:3]:  # Show top 3
                            today_warrior = next((w for w in top_warriors_local if w.get('username', '').lower() == username), None)
                            yesterday_warrior = next((w for w in yesterday_warriors if w.get('username', '').lower() == username), None)
                            if today_warrior and yesterday_warrior:
                                today_dmg = today_warrior.get('damage', 0)
                                yesterday_dmg = yesterday_warrior.get('damage', 0)
                                if yesterday_dmg > 0:
                                    change_pct = (today_dmg - yesterday_dmg) / yesterday_dmg * 100.0
                                    arrow = "▲" if change_pct > 0 else ("▼" if change_pct < 0 else "→")
                                    document.add_paragraph(f"• {username.title()}: {arrow} {change_pct:+.1f}% ({today_dmg:,} vs {yesterday_dmg:,})")
                    
                    # New warriors in top list
                    new_warriors = today_usernames - yesterday_usernames
                    if new_warriors:
                        document.add_paragraph("New warriors in top list:")
                        for username in list(new_warriors)[:3]:  # Show top 3
                            warrior = next((w for w in top_warriors_local if w.get('username', '').lower() == username), None)
                            if warrior:
                                damage = warrior.get('damage', 0)
                                country = warrior.get('country', 'Unknown')
                                document.add_paragraph(f"• {username.title()} ({country}): {damage:,} damage")

    # Economic summary (tables)
    if sections.get('economic', True):
        document.add_heading("💰 Economic Pulse", level=1)
        economic_data = summary_data.get('economic_summary', {})
        if not economic_data:
            document.add_paragraph("No economic data (endpoint may be unavailable).")
        else:
            # Top 5 best job offers worldwide
            job_offers = economic_data.get('job_offers', [])
            if job_offers:
                document.add_heading("💼 Top 5 Best Job Offers Worldwide", level=2)
                document.add_paragraph("Best paid job offers across all countries (salary in GOLD):")

                # Table without Position column
                table = document.add_table(rows=1, cols=4)
                hdr = table.rows[0].cells
                hdr[0].text = "Country"
                hdr[1].text = "Company"
                hdr[2].text = "Company ID"
                hdr[3].text = "Salary (GOLD)"

                for i, job in enumerate(job_offers[:5], 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = job.get('country', 'Unknown')
                    # Company name may not be present; leave blank or '—'
                    company_name = job.get('business_name') or job.get('company_name') or '—'
                    row_cells[1].text = str(company_name)
                    row_cells[2].text = str(job.get('business_id', 'N/A'))
                    salary_gold = job.get('salary', 0) or 0
                    # Flag clearly unrealistic salaries with an asterisk
                    note = "*" if isinstance(salary_gold, (int, float)) and salary_gold > 5 else ""
                    row_cells[3].text = f"{float(salary_gold):.6f}{note}"
                document.add_paragraph("* flagged as potentially anomalous (re-check conversion)")
                
                document.add_paragraph("")  # Dodaj pusty wiersz dla lepszego formatowania
        currency_rates = (economic_data.get('currency_rates') or {})
        codes_map = currency_codes_map or {}
        # Render currency rates as a compact table with name and rate only
        if currency_rates:
            document.add_heading("Currency rates vs GOLD", level=2)
            try:
                # Try to find the nearest previous day with currency data
                yesterday_rates = {}
                try:
                    # First try yesterday
                    if yesterday_key in historical_data:
                        yesterday_rates = (historical_data[yesterday_key].get('economic_summary') or {}).get('currency_rates') or {}
                        print(f"DEBUG: Found yesterday rates: {bool(yesterday_rates)}")
                    
                    # If no yesterday data, look for the nearest previous day
                    if not yesterday_rates and historical_data:
                        # Sort dates and find the most recent one before today
                        available_dates = sorted([d for d in historical_data.keys() if d < yesterday_key], reverse=True)
                        for date in available_dates:
                            rates = (historical_data[date].get('economic_summary') or {}).get('currency_rates') or {}
                            if rates:
                                yesterday_rates = rates
                                print(f"DEBUG: Found rates from {date}: {bool(yesterday_rates)}")
                                break
                    
                    if yesterday_rates:
                        print(f"DEBUG: Yesterday rates keys: {list(yesterday_rates.keys())[:5]}")
                except Exception as e:
                    print(f"DEBUG: Error getting yesterday rates: {e}")
                    yesterday_rates = {}
                has_diff = bool(yesterday_rates)
                print(f"DEBUG: has_diff = {has_diff}, cols = {4 if has_diff else 2}")
                cols = 4 if has_diff else 2
                table = document.add_table(rows=1, cols=cols)
                hdr = table.rows[0].cells
                hdr[0].text = "Currency"
                hdr[1].text = "Rate (GOLD per 1)"
                if has_diff:
                    hdr[2].text = "Yesterday Rate"
                    hdr[3].text = "Change %"
                # Build canonical rows
                rows_list = []
                for k, rate in currency_rates.items():
                    try:
                        cid = int(k)
                    except Exception:
                        cid = k
                    # Prefer short code; fall back to long name; finally to id
                    if codes_map:
                        name = codes_map.get(cid)
                    else:
                        name = None
                    if not name and isinstance(currencies_map, dict):
                        name = currencies_map.get(cid)
                    if not name:
                        name = str(cid)
                    try:
                        rate_val = float(rate)
                    except Exception:
                        rate_val = None
                    rows_list.append((cid, name, rate_val))
                # Sort by name
                rows_list.sort(key=lambda r: (str(r[1]).lower()))
                for cid, name, rate_val in rows_list:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(name)
                    row_cells[1].text = (f"{rate_val:.6f}" if isinstance(rate_val, (int, float)) else "no data")
                    if has_diff:
                        try:
                            prev = yesterday_rates.get(cid, yesterday_rates.get(str(cid)))
                            prev_val = float(prev) if prev is not None else None
                        except Exception:
                            prev_val = None
                        if isinstance(rate_val, (int, float)) and isinstance(prev_val, (int, float)) and prev_val != 0:
                            diff_pct = (rate_val - prev_val) / prev_val * 100.0
                            arrow = "▲" if diff_pct > 0 else ("▼" if diff_pct < 0 else "→")
                            row_cells[2].text = f"{prev_val:.6f}"
                            row_cells[3].text = f"{arrow} {diff_pct:+.2f}%"
                        else:
                            row_cells[2].text = "—"
                            row_cells[3].text = "—"
            except Exception:
                # Fallback to bullet list if table rendering fails
                for key, rate in currency_rates.items():
                    try:
                        rate_val = float(rate)
                        document.add_paragraph(f"{key}: {rate_val:.6f} GOLD", style='List Bullet')
                    except Exception:
                        document.add_paragraph(f"{key}: no data", style='List Bullet')

        cheapest_by_item = {int(k): v for k, v in (economic_data.get('cheapest_by_item') or {}).items()}
        if cheapest_by_item:
            document.add_heading("Goods — cheapest offers", level=2)
            # Helpers to parse base name and quality
            import re
            def parse_base_and_quality(name: str):
                n = (str(name) or "").strip().lower()
                m = re.search(r"\bq(\d+)\b", n)
                q = None
                if m:
                    try:
                        q = int(m.group(1))
                    except Exception:
                        q = None
                base = re.sub(r"\bq\d+\b", "", n).strip()
                base = re.sub(r"\s+", " ", base)
                return base, q

            # Build per item rows with resolved names
            per_item_rows = []  # (item_id, base, quality, display_name, country, price_cur, currency_name, price_gold, amount_at_price, avg10_gold)
            for item_id, rows in cheapest_by_item.items():
                try:
                    rows_sorted = sorted(rows, key=lambda r: (r.get('price_in_gold', 0)))
                except Exception:
                    rows_sorted = rows
                row = rows_sorted[0] if rows_sorted else {}
                country = row.get('country_name') or '—'
                currency_name = row.get('currency_name') or ''
                price_cur = row.get('price_in_currency')
                price_gold = row.get('price_in_gold')
                amount_at_price = row.get('amount_at_price')
                # Prefer new 5-offer average; fallback to legacy 10-offer key
                avg10_gold = row.get('avg5_in_gold', row.get('avg10_in_gold'))
                raw_name = (items_map.get(int(item_id)) if items_map else None) or f"item {item_id}"
                base, quality = parse_base_and_quality(raw_name)
                display_name = (str(raw_name) or f"item {item_id}").strip()
                try:
                    display_name = display_name.title()
                except Exception:
                    pass
                per_item_rows.append((item_id, base, quality, display_name, country, price_cur, currency_name, price_gold, amount_at_price, avg10_gold))

            if not per_item_rows:
                document.add_paragraph("No goods data available.")
                # Fallback to show aggregated cheapest items across countries if present later
                # (Handled by the separate section below.)
            # Define category tables: Raw Materials, Weapon, Food, Aircraft, Airplane Tickets
            def add_table(title: str, rows_src):
                if not rows_src:
                    return
                document.add_heading(title, level=3)
                
                # Check if we have yesterday's data for comparison
                yesterday_economic = historical_data.get(yesterday_key, {}).get('economic_summary', {})
                yesterday_cheapest = yesterday_economic.get('cheapest_by_item', {})
                has_yesterday = bool(yesterday_cheapest)
                
                cols = (7 if has_yesterday else 6)
                tbl = document.add_table(rows=1, cols=cols)
                h = tbl.rows[0].cells
                h[0].text = "Item"
                h[1].text = "Country"
                h[2].text = "Price (currency)"
                h[3].text = "Price (GOLD)"
                h[4].text = "Qty at price"
                h[5].text = "Avg of 5 cheapest (GOLD)"
                if has_yesterday:
                    h[6].text = "Change vs yesterday"
                
                for (_iid, _base, _q, disp, country, price_cur, currency_name, price_gold, amount_at_price, avg10_gold) in rows_src:
                    c = tbl.add_row().cells
                    c[0].text = disp
                    c[1].text = str(country)
                    c[2].text = (f"{price_cur} {currency_name}" if price_cur is not None else "—")
                    c[3].text = (f"{float(price_gold):.6f}" if isinstance(price_gold, (int, float)) else "—")
                    c[4].text = (str(int(amount_at_price)) if isinstance(amount_at_price, (int, float)) else "—")
                    c[5].text = (f"{float(avg10_gold):.6f}" if isinstance(avg10_gold, (int, float)) else "—")
                    
                    # Add price change comparison
                    if has_yesterday:
                        try:
                            yesterday_item = yesterday_cheapest.get(str(_iid), [{}])[0] if yesterday_cheapest.get(str(_iid)) else {}
                            yesterday_price = yesterday_item.get('price_in_gold')
                            if yesterday_price and isinstance(price_gold, (int, float)):
                                yesterday_val = float(yesterday_price)
                                if yesterday_val != 0:
                                    change_pct = (price_gold - yesterday_val) / yesterday_val * 100.0
                                    arrow = "▲" if change_pct > 0 else ("▼" if change_pct < 0 else "→")
                                    c[6].text = f"{arrow} {change_pct:+.2f}%"
                                else:
                                    c[6].text = "—"
                            else:
                                c[6].text = "—"
                        except (ValueError, TypeError):
                            c[6].text = "—"

            # Grouping by fixed categories
            def is_raw(base: str) -> bool:
                return base in {"grain", "iron", "titanium", "fuel", "raw material", "raw materials"}

            def is_weapon(base: str) -> bool:
                return base.startswith("weapon")

            def is_food(base: str) -> bool:
                return base.startswith("food")

            def is_aircraft(base: str) -> bool:
                # Aircraft items but not tickets
                if is_air_tickets(base):
                    return False
                return (
                    base.startswith("aircraft")
                    or (base.startswith("airplane") and "ticket" not in base)
                    or (base.startswith("plane") and "ticket" not in base)
                )

            def is_air_tickets(base: str) -> bool:
                return base.startswith("airplane ticket") or base.startswith("air ticket") or ("ticket" in base and ("air" in base or "plane" in base))

            raws = [r for r in per_item_rows if is_raw(r[1])]
            weapons = [r for r in per_item_rows if is_weapon(r[1])]
            foods = [r for r in per_item_rows if is_food(r[1])]
            aircrafts = [r for r in per_item_rows if is_aircraft(r[1])]
            tickets = [r for r in per_item_rows if is_air_tickets(r[1])]

            # Render exactly five tables
            add_table("Raw Materials", sorted(raws, key=lambda r: (r[1], r[2] or 0, r[3])))
            add_table("Weapon", sorted(weapons, key=lambda r: (r[2] or 0, r[3])))
            add_table("Food", sorted(foods, key=lambda r: (r[2] or 0, r[3])))
            add_table("Aircraft", sorted(aircrafts, key=lambda r: (r[2] or 0, r[3])))
            add_table("Airplane Tickets", sorted(tickets, key=lambda r: (r[2] or 0, r[3])))

            # Removed price summary table per request
        
        # Cheapest goods from all countries
        cheapest_items_all = economic_data.get('cheapest_items_all_countries', {})
        
        # Debug logging to help identify the issue
        print(f"DEBUG: cheapest_items_all_countries found: {bool(cheapest_items_all)}")
        if cheapest_items_all:
            print(f"DEBUG: Number of items: {len(cheapest_items_all)}")
            print(f"DEBUG: Sample keys: {list(cheapest_items_all.keys())[:5] if cheapest_items_all else 'None'}")
        else:
            print(f"DEBUG: Available keys in economic_data: {list(economic_data.keys())}")
        
        if cheapest_items_all:
            document.add_heading("🌍 Cheapest Goods from All Countries", level=2)
            document.add_paragraph("Cheapest goods of each type from all countries:")
            
            # Group goods by type and quality
            import re
            
            def parse_item_type_and_quality(name: str):
                """Parse item name and return type and quality"""
                name_lower = str(name).lower()
                
                # Find quality (Q1, Q2, Q3, Q4, Q5)
                quality_match = re.search(r'\bq(\d+)\b', name_lower)
                quality = int(quality_match.group(1)) if quality_match else None
                
                # Remove quality from name and determine type
                base_name = re.sub(r'\bq\d+\b', '', name_lower).strip()
                
                if any(word in base_name for word in ["grain", "zboże"]):
                    return "Grain", quality
                elif any(word in base_name for word in ["iron"]):
                    return "Iron", quality
                elif any(word in base_name for word in ["titanium", "tytan"]):
                    return "Titanium", quality
                elif any(word in base_name for word in ["fuel", "paliwo"]):
                    return "Fuel", quality
                elif any(word in base_name for word in ["food"]):
                    return "Food", quality
                elif "weapon" in base_name or "broń" in base_name:
                    return "Weapon", quality
                elif "aircraft" in base_name or ("airplane" in base_name and "ticket" not in base_name):
                    return "Aircraft", quality
                elif "ticket" in base_name:
                    return "Airplane Ticket", quality
                else:
                    return "Other", quality
            
            # Group goods by type and quality
            grouped_items = {}
            for item_id, item_list in cheapest_items_all.items():
                try:
                    if item_list and len(item_list) > 0:
                        for item in item_list:
                            item_name = item.get('item_name', '')
                            item_type, quality = parse_item_type_and_quality(item_name)
                            
                            # Create group key
                            if quality is not None:
                                group_key = f"{item_type} Q{quality}"
                            else:
                                group_key = f"{item_type} (No Quality)"
                            
                            if group_key not in grouped_items:
                                grouped_items[group_key] = []
                            grouped_items[group_key].append(item)
                except Exception as e:
                    print(f"DEBUG: Error processing item {item_id}: {e}")
                    continue
            
            print(f"DEBUG: Grouped items: {list(grouped_items.keys())}")
            
            # Display each group in a separate table
            for group_name, items in grouped_items.items():
                if items:
                    document.add_heading(f"{group_name}", level=3)
                    
                    table = document.add_table(rows=1, cols=6)
                    hdr = table.rows[0].cells
                    hdr[0].text = "Good"
                    hdr[1].text = "Country"
                    hdr[2].text = "Price (currency)"
                    hdr[3].text = "Price (GOLD)"
                    hdr[4].text = "Quantity"
                    hdr[5].text = "Average of 5 cheapest (GOLD)"
                    
                    # Sort goods by GOLD price
                    sorted_items = sorted(items, key=lambda x: x.get('price_gold', 0))
                    
                    for item in sorted_items:
                        try:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(item.get('item_name', 'Nieznany'))
                            row_cells[1].text = str(item.get('country', 'Nieznany'))
                            row_cells[2].text = f"{item.get('price_currency', 0):.3f} {item.get('currency_name', '')}"
                            row_cells[3].text = f"{item.get('price_gold', 0):.6f}"
                            qty = item.get('amount')
                            row_cells[4].text = (str(int(qty)) if isinstance(qty, (int, float)) else "—")
                            avg5 = item.get('avg5_in_gold')
                            row_cells[5].text = (f"{float(avg5):.6f}" if isinstance(avg5, (int, float)) else "—")
                        except Exception as e:
                            print(f"DEBUG: Error processing row for item {item}: {e}")
                            continue
                    
                    document.add_paragraph("")  # Spacer for better formatting
        else:
            # Add fallback section if no cheapest items data
            document.add_heading("🌍 Cheapest Goods from All Countries", level=2)
            document.add_paragraph("No data on cheapest goods from all countries.")
            document.add_paragraph("Available keys in economic data:")
            for key in economic_data.keys():
                document.add_paragraph(f"• {key}", style='List Bullet')

        # Production analysis as part of economic section (always show if enabled)
        production_data = summary_data.get('production_data', [])
        if sections.get('production', False) and production_data:
            document.add_heading("🏭 Regional Productivity Analysis", level=1)
            document.add_paragraph("Analysis of production efficiency in different regions for various goods.")
            
            # Group by item type instead of efficiency score
            items_groups = {}
            for data in production_data:
                # Handle both object and dictionary formats
                if hasattr(data, 'item_name'):
                    item_name = data.item_name
                else:
                    item_name = data.get('item_name', 'Unknown')
                
                if item_name not in items_groups:
                    items_groups[item_name] = []
                items_groups[item_name].append(data)
            
            for item_name, items in items_groups.items():
                if items:
                    # Sort by efficiency score (descending)
                    items.sort(key=lambda x: x.efficiency_score if hasattr(x, 'efficiency_score') else x.get('efficiency_score', 0), reverse=True)
                    
                    document.add_heading(f"Product: {item_name.title()}", level=2)
                    
                    # Create table with all quality levels and NPC wages
                    table = document.add_table(rows=1, cols=11)
                    hdr = table.rows[0].cells
                    hdr[0].text = "Region"
                    hdr[1].text = "Country"
                    hdr[2].text = "Score"
                    hdr[3].text = "Regional Bonus"
                    hdr[4].text = "Country Bonus"
                    hdr[5].text = "Q1"
                    hdr[6].text = "Q2"
                    hdr[7].text = "Q3"
                    hdr[8].text = "Q4"
                    hdr[9].text = "Q5"
                    hdr[10].text = "NPC Wages (GOLD)"
                    
                    # Show top 10 from each item type
                    for item in items[:10]:
                        row_cells = table.add_row().cells
                        # Handle both object and dictionary formats
                        if hasattr(item, 'region_name'):
                            row_cells[0].text = item.region_name
                            row_cells[1].text = item.country_name
                            row_cells[2].text = f"{item.efficiency_score:.2f}"
                            row_cells[3].text = f"{item.regional_bonus:.1%}"
                            row_cells[4].text = f"{item.country_bonus:.1f}%"
                            row_cells[5].text = str(item.production_q1)
                            row_cells[6].text = str(item.production_q2)
                            row_cells[7].text = str(item.production_q3)
                            row_cells[8].text = str(item.production_q4)
                            row_cells[9].text = str(item.production_q5)
                            row_cells[10].text = f"{item.npc_wages:.2f}"
                        else:
                            row_cells[0].text = item.get('region_name', 'Unknown')
                            row_cells[1].text = item.get('country_name', 'Unknown')
                            row_cells[2].text = f"{item.get('efficiency_score', 0):.2f}"
                            row_cells[3].text = f"{item.get('regional_bonus', 0):.1%}"
                            row_cells[4].text = f"{item.get('country_bonus', 0):.1f}%"
                            row_cells[5].text = str(item.get('production_q1', 0))
                            row_cells[6].text = str(item.get('production_q2', 0))
                            row_cells[7].text = str(item.get('production_q3', 0))
                            row_cells[8].text = str(item.get('production_q4', 0))
                            row_cells[9].text = str(item.get('production_q5', 0))
                            row_cells[10].text = f"{item.get('npc_wages', 0):.2f}"
                    
                    document.add_paragraph("")  # Spacer

        # Add detailed production analysis by item type (similar to HTML)

    # Create reports directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    file_name = f"daily_report_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
    file_path = os.path.join(output_dir, file_name)
    document.save(file_path)
    print(f"DOCX report successfully generated as '{file_path}'")
    return file_path


