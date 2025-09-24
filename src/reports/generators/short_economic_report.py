#!/usr/bin/env python3
"""
Short economic report - contains all currency rates,
cheapest item of each type and best region for each product.
"""

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.services.economy_service import (
    fetch_cheapest_items_from_all_countries,
    fetch_items_by_type,
    fetch_best_jobs_from_all_countries
)
from src.data.repositories.sqlite_repository import SQLiteCurrencyRepository
from src.reports.generators.production_report import ProductionAnalyzer
from src.data.storage.cache import load_historical_data
from datetime import timedelta


def generate_short_economic_report(
    output_dir: str = "reports",
    sections: dict = None
) -> str:
    """
    Generates a short economic report in DOCX format.
    
    Args:
        output_dir: Directory to save the report
        sections: Sections to include (all by default)
    
    Returns:
        Path to the generated file
    """
    
    # Default sections
    if sections is None:
        sections = {
            'currency_rates': True,
            'cheapest_items': True,
            'best_regions': True,
            'highest_wages': True
        }
    
    print("🚀 Generating short economic report...")
    
    # Load historical data for currency rate comparison
    historical_data = load_historical_data()
    
    # Get economic data using refactored service (from database)
    print("📊 Fetching economic data from database...")
    try:
        # Initialize repositories and service
        currency_repo = SQLiteCurrencyRepository("data/eclesiar.db")
        
        # Create minimal dependencies for economy service
        deps = ServiceDependencies(
            country_repo=None,
            currency_repo=currency_repo,
            region_repo=None,
            item_repo=None,
            market_repo=None,
            production_repo=None,
            report_repo=None
        )
        economy_service = EconomyService(deps)
        
        # First try to get from database
        eco_countries, currencies_map, currency_codes_map, gold_id = economy_service.get_countries_and_currencies()
        
        # If no data in database, fall back to API
        if not eco_countries or not currencies_map:
            print("📊 No data in database, falling back to API...")
            from src.core.services.economy_service import fetch_countries_and_currencies
            eco_countries, currencies_map, currency_codes_map, gold_id = fetch_countries_and_currencies()
        
        if not eco_countries or not currencies_map:
            print("❌ Error: Cannot fetch economic data")
            return None
        
        # Get currency rates from database first, then API if needed
        print("💰 Fetching currency rates from database...")
        currency_rates = economy_service.build_currency_rates_map(currencies_map, gold_id)
        
        # If no currency rates from database, fall back to API
        if not currency_rates:
            print("💰 No currency rates in database, falling back to API...")
            from src.core.services.economy_service import build_currency_rates_map
            currency_rates = build_currency_rates_map(currencies_map, gold_id)
        
    except Exception as e:
        print(f"❌ Error with database service: {e}")
        print("📊 Falling back to API...")
        from src.core.services.economy_service import fetch_countries_and_currencies, build_currency_rates_map
        eco_countries, currencies_map, currency_codes_map, gold_id = fetch_countries_and_currencies()
        currency_rates = build_currency_rates_map(currencies_map, gold_id)
    
    # Get items map
    print("📦 Fetching items map...")
    items_map = fetch_items_by_type("economic")
    
    # Get cheapest items
    print("🛒 Fetching cheapest items...")
    cheapest_items = fetch_cheapest_items_from_all_countries(
        eco_countries, items_map, currency_rates, gold_id
    )
    
    # Get best jobs (highest wages)
    print("💼 Fetching best job offers...")
    best_jobs = fetch_best_jobs_from_all_countries(
        eco_countries, currency_rates, gold_id
    )
    
    # Get region data for productivity analysis
    print("🏭 Fetching region data...")
    from src.core.services.regions_service import fetch_and_process_regions
    regions_data, regions_summary = fetch_and_process_regions(eco_countries)
    
    # Create DOCX document
    document = Document()
    
    # Title
    title = document.add_heading("Short Economic Report Eclesiar", level=1)
    for run in title.runs:
        run.font.size = Pt(24)
    
    # Date information
    document.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("")
    
    # 1. Currency rates
    if sections.get('currency_rates', True):
        print("📈 Adding currency rates section...")
        document.add_heading("💰 Currency Rates vs GOLD", level=1)
        
        if currency_rates:
            # Get yesterday's rates for comparison
            yesterday_key = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
            yesterday_rates = {}
            
            try:
                if yesterday_key in historical_data:
                    yesterday_rates = (historical_data[yesterday_key].get('economic_summary') or {}).get('currency_rates') or {}
                    print(f"DEBUG: Found yesterday rates: {bool(yesterday_rates)}")
                
                # If no yesterday data, look for the nearest previous day
                if not yesterday_rates and historical_data:
                    available_dates = sorted([d for d in historical_data.keys() if d < yesterday_key], reverse=True)
                    for date in available_dates:
                        rates = (historical_data[date].get('economic_summary') or {}).get('currency_rates') or {}
                        if rates:
                            yesterday_rates = rates
                            print(f"DEBUG: Found rates from {date}: {bool(yesterday_rates)}")
                            break
            except Exception as e:
                print(f"DEBUG: Error getting yesterday rates: {e}")
                yesterday_rates = {}
            
            # Sort rates alphabetically by currency name
            sorted_rates = []
            for currency_id, rate in currency_rates.items():
                currency_name = currencies_map.get(currency_id, f"Currency {currency_id}")
                currency_code = currency_codes_map.get(currency_id, "")
                display_name = currency_code if currency_code else currency_name
                sorted_rates.append((display_name, rate, currency_name, currency_id))
            
            sorted_rates.sort(key=lambda x: x[0].lower())
            
            # Determine table columns based on whether we have historical data
            has_historical = bool(yesterday_rates)
            cols = 5 if has_historical else 3
            
            # Currency rates table
            table = document.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Currency"
            hdr_cells[1].text = "Code"
            hdr_cells[2].text = "Rate (GOLD per 1)"
            if has_historical:
                hdr_cells[3].text = "Previous Rate"
                hdr_cells[4].text = "Change %"
            
            # Data rows
            for display_name, rate, full_name, currency_id in sorted_rates:
                row_cells = table.add_row().cells
                row_cells[0].text = full_name
                row_cells[1].text = display_name
                row_cells[2].text = f"{rate:.6f}"
                
                if has_historical:
                    # Get previous rate for this currency
                    prev_rate = yesterday_rates.get(str(currency_id))
                    if prev_rate is not None:
                        row_cells[3].text = f"{prev_rate:.6f}"
                        # Calculate percentage change
                        if prev_rate > 0:
                            change_pct = ((rate - prev_rate) / prev_rate) * 100
                            row_cells[4].text = f"{change_pct:+.2f}%"
                        else:
                            row_cells[4].text = "N/A"
                    else:
                        row_cells[3].text = "N/A"
                        row_cells[4].text = "N/A"
        else:
            document.add_paragraph("No currency rate data available.")
        
        document.add_paragraph("")
    
    # 2. Cheapest items of each type
    if sections.get('cheapest_items', True):
        print("🛒 Adding cheapest items section...")
        document.add_heading("🛒 Cheapest Items of Each Type", level=1)
        
        if cheapest_items:
            # Group items by type and quality
            grouped_items = _group_items_by_type(cheapest_items)
            
            for product_type, quality_items in grouped_items.items():
                if quality_items:
                    document.add_heading(f"{product_type}", level=2)
                    
                    # Create table with all quality levels
                    table = document.add_table(rows=1, cols=7)
                    table.style = 'Table Grid'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Quality"
                    hdr_cells[1].text = "Item"
                    hdr_cells[2].text = "Country"
                    hdr_cells[3].text = "Price (currency)"
                    hdr_cells[4].text = "Price (GOLD)"
                    hdr_cells[5].text = "Quantity"
                    hdr_cells[6].text = "Average of 5 cheapest (GOLD)"
                    
                    # Sort quality levels (Q1, Q2, Q3, Q4, Q5)
                    sorted_qualities = sorted(quality_items.keys(), key=lambda x: int(x[1:]) if x[1:].isdigit() else 0)
                    
                    for quality_level in sorted_qualities:
                        items = quality_items[quality_level]
                        if items:
                            # Find cheapest item of this quality level
                            cheapest_item = min(items, key=lambda x: x.get('price_gold', float('inf')))
                            
                            # Data row with bold formatting for the cheapest item
                            row_cells = table.add_row().cells
                            
                            # Quality level
                            row_cells[0].text = quality_level
                            
                            # Make the item name bold
                            item_name_cell = row_cells[1]
                            item_name_cell.text = str(cheapest_item.get('item_name', 'Unknown'))
                            for paragraph in item_name_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                            
                            # Make the country name bold
                            country_cell = row_cells[2]
                            country_cell.text = str(cheapest_item.get('country', 'Unknown'))
                            for paragraph in country_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                            
                            # Regular formatting for price in currency
                            row_cells[3].text = f"{cheapest_item.get('price_currency', 0):.3f} {cheapest_item.get('currency_name', '')}"
                            
                            # Make the price in GOLD bold (most important)
                            price_gold_cell = row_cells[4]
                            price_gold_cell.text = f"{cheapest_item.get('price_gold', 0):.6f}"
                            for paragraph in price_gold_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                            
                            # Regular formatting for other cells
                            qty = cheapest_item.get('amount')
                            row_cells[5].text = (str(int(qty)) if isinstance(qty, (int, float)) else "—")
                            avg5 = cheapest_item.get('avg5_in_gold')
                            row_cells[6].text = (f"{float(avg5):.6f}" if isinstance(avg5, (int, float)) else "—")
                    
                    document.add_paragraph("")
        else:
            document.add_paragraph("No cheapest items data available.")
        
        document.add_paragraph("")
    
    # 3. Production examples for each product (Q1-Q5)
    if sections.get('best_regions', True) and regions_data:
        print("🏭 Adding production examples section...")
        document.add_heading("🏭 Production Examples by Product and Quality", level=1)
        
        # Analyze productivity
        analyzer = ProductionAnalyzer()
        production_data = analyzer.analyze_all_regions(regions_data)
        
        if production_data:
            # Group by product type
            products_groups = {}
            for data in production_data:
                item_name = data.item_name if hasattr(data, 'item_name') else data.get('item_name', 'Unknown')
                if item_name not in products_groups:
                    products_groups[item_name] = []
                products_groups[item_name].append(data)
            
            # For each product show one example of each quality (Q1-Q5)
            for product_name, items in products_groups.items():
                if items:
                    # Sort by efficiency_score
                    items.sort(
                        key=lambda x: x.efficiency_score if hasattr(x, 'efficiency_score') else x.get('efficiency_score', 0), 
                        reverse=True
                    )
                    
                    document.add_heading(f"Product: {product_name.title()}", level=2)
                    
                    # Table with one example of each quality
                    table = document.add_table(rows=1, cols=10)
                    table.style = 'Table Grid'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Region"
                    hdr_cells[1].text = "Country"
                    hdr_cells[2].text = "Score"
                    hdr_cells[3].text = "Regional Bonus"
                    hdr_cells[4].text = "Country Bonus"
                    hdr_cells[5].text = "Q1"
                    hdr_cells[6].text = "Q2"
                    hdr_cells[7].text = "Q3"
                    hdr_cells[8].text = "Q4"
                    hdr_cells[9].text = "Q5"
                    
                    # Take the best region (it has all Q1-Q5 values)
                    best_region = items[0]
                    
                    # Data row
                    row_cells = table.add_row().cells
                    if hasattr(best_region, 'region_name'):
                        row_cells[0].text = best_region.region_name
                        row_cells[1].text = best_region.country_name
                        row_cells[2].text = f"{best_region.efficiency_score:.2f}"
                        row_cells[3].text = f"{best_region.regional_bonus:.1%}"
                        row_cells[4].text = f"{best_region.country_bonus:.1f}%"
                        row_cells[5].text = str(best_region.production_q1)
                        row_cells[6].text = str(best_region.production_q2)
                        row_cells[7].text = str(best_region.production_q3)
                        row_cells[8].text = str(best_region.production_q4)
                        row_cells[9].text = str(best_region.production_q5)
                    else:
                        row_cells[0].text = best_region.get('region_name', 'Unknown')
                        row_cells[1].text = best_region.get('country_name', 'Unknown')
                        row_cells[2].text = f"{best_region.get('efficiency_score', 0):.2f}"
                        row_cells[3].text = f"{best_region.get('regional_bonus', 0):.1%}"
                        row_cells[4].text = f"{best_region.get('country_bonus', 0):.1f}%"
                        row_cells[5].text = str(best_region.get('production_q1', 0))
                        row_cells[6].text = str(best_region.get('production_q2', 0))
                        row_cells[7].text = str(best_region.get('production_q3', 0))
                        row_cells[8].text = str(best_region.get('production_q4', 0))
                        row_cells[9].text = str(best_region.get('production_q5', 0))
                    
                    document.add_paragraph("")
        else:
            document.add_paragraph("No region productivity data available.")
    
    # 4. Top 10 Highest Wages
    if sections.get('highest_wages', True) and best_jobs:
        print("💼 Adding highest wages section...")
        document.add_heading("💼 Top 10 Highest Salaries Worldwide", level=1)
        
        # Add disclaimer about data accuracy
        disclaimer = document.add_paragraph()
        disclaimer.add_run("⚠️ Note: ").bold = True
        disclaimer.add_run("Job market data is fetched from API and may include offers that are no longer available in-game due to being filled or expired. Some very high salaries (like 62 HUF from Hungary) may be outdated. Actual highest salaries may differ from what is currently visible in the game.")
        
        # Extend the list to show top 10 instead of top 5
        all_jobs = fetch_best_jobs_from_all_countries(eco_countries, currency_rates, gold_id)
        # Sort by salary_gold descending and take top 10
        all_jobs.sort(key=lambda x: x.get("salary_gold", 0), reverse=True)
        top_10_jobs = all_jobs[:10]
        
        if top_10_jobs:
            # Create table for highest salaries
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Rank"
            hdr_cells[1].text = "Country"
            hdr_cells[2].text = "Business ID"
            hdr_cells[3].text = "Salary (GOLD)"
            hdr_cells[4].text = "Salary (Local)"
            
            # Data rows
            for i, job in enumerate(top_10_jobs, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = job.get('country_name', 'Unknown')
                # Użyj business_id z API
                business_id = job.get('business_id') or job.get('company_id') or f"Job-{i}"
                row_cells[2].text = str(business_id)
                
                # Make salary bold (most important column)
                salary_cell = row_cells[3]
                salary_gold = job.get('salary_gold', 0)
                salary_cell.text = f"{salary_gold:.6f}"
                for paragraph in salary_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                
                # Local salary with currency
                salary_original = job.get('salary_original', 0)
                currency_name = job.get('currency_name', 'N/A')
                row_cells[4].text = f"{salary_original:.6f} {currency_name}"
        else:
            document.add_paragraph("No job offers data available.")
        
        document.add_paragraph("")
    
    # Save document
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
    filename = f"short_economic_report_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    document.save(filepath)
    print(f"✅ Short economic report generated: {filepath}")
    
    return filepath


def _group_items_by_type(cheapest_items: Dict[int, List[Dict[str, Any]]]) -> Dict[str, Dict[str, List[Dict[str, Any]]]]:
    """
    Groups items by type (weapon, iron, etc.) and then by quality level (Q1, Q2, Q3, Q4, Q5)
    Returns: {product_type: {quality_level: [items]}}
    """
    import re
    
    def parse_item_type(name: str) -> str:
        """Determines item type based on name"""
        name_lower = str(name).lower()
        
        if any(word in name_lower for word in ["grain", "zboże"]):
            return "Grain"
        elif any(word in name_lower for word in ["iron"]):
            return "Iron"
        elif any(word in name_lower for word in ["titanium", "tytan"]):
            return "Titanium"
        elif any(word in name_lower for word in ["fuel", "paliwo"]):
            return "Fuel"
        elif any(word in name_lower for word in ["food"]):
            return "Food"
        elif "weapon" in name_lower or "broń" in name_lower:
            return "Weapon"
        elif "aircraft" in name_lower or ("airplane" in name_lower and "ticket" not in name_lower):
            return "Aircraft"
        elif "ticket" in name_lower:
            return "Airplane Ticket"
        else:
            return "Other"
    
    def parse_item_quality(name: str) -> str:
        """Determines item quality level based on name"""
        name_lower = str(name).lower()
        
        # Look for quality patterns like "q1", "q2", etc.
        quality_match = re.search(r'q(\d+)', name_lower)
        if quality_match:
            quality_num = int(quality_match.group(1))
            if 1 <= quality_num <= 5:
                return f"Q{quality_num}"
        
        # If no quality found, assume Q1
        return "Q1"
    
    grouped = {}
    
    for item_id, item_list in cheapest_items.items():
        if item_list and len(item_list) > 0:
            for item in item_list:
                item_name = item.get('item_name', '')
                item_type = parse_item_type(item_name)
                item_quality = parse_item_quality(item_name)
                
                if item_type not in grouped:
                    grouped[item_type] = {}
                if item_quality not in grouped[item_type]:
                    grouped[item_type][item_quality] = []
                grouped[item_type][item_quality].append(item)
    
    return grouped


def main():
    """Main function for testing"""
    try:
        report_path = generate_short_economic_report()
        if report_path:
            print(f"✅ Report generated successfully: {report_path}")
        else:
            print("❌ Error generating report")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
